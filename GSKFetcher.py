import logging
import os
from MyTools.fetcher import Fetcher
from GSKParser import GSKPageParser
from GSKParser import GSKPaperParser
from MyTools.tools import URLTool
import urllib.parse
import urllib.request
from docx import Document


os.environ["NLS_LANG"] = "SIMPLIFIED CHINESE_CHINA.UTF8"

logFileHandler = logging.FileHandler(
    filename="GSK.log",
    mode="w",
    encoding="utf-8"
)
console = logging.StreamHandler()

formatter = logging.Formatter("%(levelname)s %(asctime)s %(lineno)d: %(message)s")
console.setFormatter(formatter)
logFileHandler.setFormatter(formatter)

logging.getLogger().addHandler(logFileHandler)
logging.getLogger().addHandler(console)

console.setLevel(logging.INFO)
logging.getLogger().setLevel(logging.DEBUG)


class GSKSpider:
    ROOT_URL = "https://health.gsk-china.com/service-and-tools/elsevier"
    __DEFAULT_PAGE = "/jacc/volume-67,issue-7,february-23-2016/"
    __COOKIE = "ta=KUCFTiQ%2brx4Hu7ba1eSh9w%3d%3d; taid=1112; __utmt=1; notification=1; announce=1; stoken=Ljagj" \
               "iORONckpy9fhJZc9sKbiP9up3S0rVGLFcIAZGbSaet6GvWs4ghAEsupdqkXrSw0Ioh245WZI3/6KMCZT4mUDnmX4p6Yg/MW7" \
               "uSPlNxtM3pzILGZ5Dp2WCIeWI06bHgYS5fqIAcqOTEpLJRwmHaWxGCAQVjIDNVIc/y7PzkfmdVP62f8Q1wmmZ/mnZqKy37FJ" \
               "t0FR+G6qitMh0Nf0RUjP9z9jwC2bONTfCCr65lkP649SbOfr/ir9IdIC+vXFt5F6W2b1HHwMulkqSBfMU8F9vAlKP5RGdUkg" \
               "dy39CNWx2IML4j9tb+h1vczYQeA6GcSORcNkwgVX8Z9rPIETgU11fLg7hQSJEfB84yoWVjqOHap8ZsANFVuv31zLuc8hahis" \
               "5r3ZKDR87HcgXF5QPu3Yi8h4CaK4TkRzd2QmOhGMVXHTdUjp2+kBFWQZ1jy77hQMSNVK5Xks1nBekQgHYoSZDy+IBzEOaWWJ" \
               "rvGDCPavF/SUjh8fpBoRqS2jJ2ZE7tzxgNA5DcrINVF6snBSA==; __utma=159908132.374290233.1464178012.14641" \
               "78012.1464178012.1; __utmb=159908132.16.9.1464178570013; __utmc=159908132; __utmz=159908132.1464" \
               "178012.1.1.utmcsr=(direct)|utmccn=(direct)|utmcmd=(none); yourAuthCookie=FEC7F8F85E66FD822E2AE95" \
               "F85772C8413239A460FFD79B534D11306CD7BB18FD435EBB86FEE8E49F6AF36D9161B6567F2CE9F8C02EA6879B1054F5" \
               "3E0F6972C66A26C6C29BE9B3D8D79F986ADD440CB67CEED0EB3B94AB7CD89E38C9EB4DA076A68A36EC9D6B718529066A" \
               "FDFAB70D1597685060869B33951E2DC630707734B40E6E9223B8A6ADEFE2C337326A7D0040837BDA21B6815D3F7E82E0" \
               "FF559C1CF; yourAccount=Nmh5ZrGLdpkgl/8WCZoFGBI0anshyjDEzaDKxl+LCoo=; keyEncryption=a9rFGCh9G9IxN" \
               "z0iL8bp6A%3d%3d"

    __HEADER = {
        ("accept", "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8"),
        ("accept-encoding", "gzip,deflate"),
        ("accept-language", "zh-CN,zh;q=0.8"),
        ("cache-control", "no-cache"),
        ("connection", "keep-alive"),
        ("cookie", __COOKIE),
        ("content-type", "application/x-www-form-urlencoded"),
        ("host", "health.gsk-china.com"),
        ("Pragma", "no-cache"),
        ("referer", "https://health.gsk-china.com/Login/?url=%2fservice-and-tools%2felsevier%2fjacc%2fvolume-67%2cissue"
                    "-7%2cfebruary-23-2016%2f&d=LoginRequest"),
        ("user-agent", "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) "
                       "Chrome/38.0.2125.122 Safari/537.36 SE 2.X MetaSr 1.0"),
    }

    def __init__(self):
        self._opener = URLTool.get_url_opener(
            header=self.__HEADER,
            need_cookie=True
        )
        with open("StartPoint", "r") as sp:
            self.start_pages = [url.strip() for url in sp]
        with open("Visited", "r") as v:
            self.visted_pages = set(url.strip() for url in v)

    @staticmethod
    def _fetch_all_info(fetcher, parser, method_name, list_method="extend"):
        all_results_list = []
        while fetcher.taskleft():
            current_url, page_content = fetcher.pop()
            if page_content == "":
                continue
            parser.feed(page_content)
            method = getattr(parser, method_name)
            current_list = method()
            if current_list is not None and len(current_list) > 0:
                if list_method == "extend":
                    all_results_list.extend(current_list)
                elif list_method == "append":
                    all_results_list.append(current_list)
                else:
                    all_results_list.extend(current_list)
        return all_results_list

    def _get_journal_url_list(self, main_page_url):
        fetcher = Fetcher(
            threads_num=1,
            interval=0.5,
            opener=self._opener
        )
        fetcher.push(main_page_url)
        return self._fetch_all_info(fetcher, GSKPageParser(), "get_journals")

    def _get_issue_url_list(self, journal):
        journal_url = urllib.request.quote(journal, safe=":/%")
        fetcher = Fetcher(
            threads_num=1,
            interval=0.5,
            opener=self._opener
        )
        fetcher.push(journal_url)
        return self._fetch_all_info(fetcher, GSKPageParser(), "get_issues")

    def _get_paper_url_list(self, issue):
        issue_url = urllib.request.quote(issue, safe=":/%")
        fetcher = Fetcher(
            threads_num=1,
            interval=0.5,
            opener=self._opener
        )
        fetcher.push(issue_url)
        return self._fetch_all_info(fetcher, GSKPageParser(), "get_papers")

    def _get_abstracts(self, paper_list):
        fetcher = Fetcher(
            threads_num=10,
            interval=2,
            opener=self._opener
        )
        for paper in paper_list:
            logging.info("  Getting abstract from %s..." % paper)
            paper_url = urllib.request.quote(paper, safe=":/%")
            fetcher.push(paper_url)
        return self._fetch_all_info(fetcher, GSKPaperParser(), "get_abstract", list_method="append")

    @staticmethod
    def _save_results(filename, abstracts):
        document = Document()
        for title, content in abstracts:
            document.add_heading(title, level=0)
            for p in content.split("\r\n"):
                document.add_paragraph(p)
            document.add_page_break()
        document.save("./result/" + filename)

    def start(self):
        logging.info("Spider started, trying to get journal urls from page.")
        journal_urls = self.start_pages
        logging.info("Got %d journals." % len(journal_urls))

        new_start_page = []
        for journal in journal_urls:
            logging.info("Trying to get issue url of Journal: %s." % journal)
            issue_urls = self._get_issue_url_list(journal)
            logging.info("Got %d issues." % len(issue_urls))

            if issue_urls:
                new_start_page.append(issue_urls[0])

            for issue in issue_urls:
                if issue in self.visted_pages:
                    logging.info("Already visted %s before, will skip" % issue)
                    continue
                self.visted_pages.add(issue)

                url_array = issue.split("/")
                journal_name = url_array[-3]
                issue_name = url_array[-2]

                logging.info("Trying to get paper url in %s %s." % (journal_name, issue_name))
                paper_urls = self._get_paper_url_list(issue)
                logging.info("Got %d papers." % len(paper_urls))

                logging.info("Trying to get abstracts of All Papers of issue %s." % issue_name)
                abstracts = self._get_abstracts(paper_urls)
                logging.info("Got %d abstracts from issue %s." % (len(abstracts), issue_name))

                filename = "%s_%s.docx" % (journal_name, issue_name)
                logging.info("Saving abstracts of papers into %s..." % filename)
                self._save_results(filename, abstracts)
                logging.info("%d abstracts saved." % len(abstracts))

        logging.info("Updating StartPoint and Visted...")
        with open("StartPoint", "w") as sp:
            for url in new_start_page:
                sp.write("%s" % url)
        with open("Visited", "w") as v:
            for url in self.visted_pages:
                v.write("%s" % url)

if __name__ == "__main__":
    spider = GSKSpider()
    spider.start()

